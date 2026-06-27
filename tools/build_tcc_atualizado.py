from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "tcc_atualizado"
OUT_DOCX = OUT_DIR / "TCC_TechOS_Flow_Mateus_Lima_atualizado_Railway.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def add_paragraph(doc, text="", style=None, align=None, bold=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run(item)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        header_cells[i].text = h
        set_cell_shading(header_cells[i], "E8EEF5")
        for p in header_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 25 else WD_ALIGN_PARAGRAPH.CENTER
    set_table_widths(table, widths_cm)
    doc.add_paragraph()
    return table


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after in [
        ("Heading 1", 14, 18, 8),
        ("Heading 2", 13, 14, 6),
        ("Heading 3", 12, 10, 4),
    ]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.5

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(12)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_after = Pt(4)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_doc(doc)

    # Capa
    for text in [
        "CENTRO UNIVERSITÁRIO UNINORTE",
        "CURSO DE BACHARELADO DE SISTEMAS DE INFORMAÇÃO",
    ]:
        add_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    for _ in range(6):
        doc.add_paragraph()
    add_paragraph(doc, "MATEUS LIMA DA SILVA BRAGA", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    for _ in range(5):
        doc.add_paragraph()
    add_paragraph(
        doc,
        "TechOS Flow: Sistema Web para Gestão Digital de Ordens de Serviço no Setor de Saneamento",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    for _ in range(6):
        doc.add_paragraph()
    p = add_paragraph(
        doc,
        "Trabalho apresentado ao curso de graduação em Sistemas de Informação do Centro Universitário Uninorte (União Educacional do Norte), como requisito acadêmico de conclusão de curso.\n\nProfessor orientador: Rodrigo Garcia",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    p.paragraph_format.left_indent = Cm(7)
    for _ in range(4):
        doc.add_paragraph()
    add_paragraph(doc, "RIO BRANCO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "2026", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_page_break()

    footer = doc.sections[0].footer.paragraphs[0]
    add_page_number(footer)

    # Resumo
    add_paragraph(doc, "RESUMO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_body(
        doc,
        "O presente trabalho apresenta a atualização do desenvolvimento do sistema web TechOS Flow, uma solução voltada à abertura, execução, acompanhamento e análise de ordens de serviço no contexto operacional do saneamento. A versão consolidada do projeto deixou de ser apenas uma proposta de automação e passou a contemplar um sistema funcional com autenticação por perfis, ciclo completo de OS, evidências com anexos e geolocalização, relatórios administrativos, exportações, cálculo de horas extras, banco de folgas, recuperação de senha por e-mail e preparação para implantação em ambiente de nuvem. A pesquisa caracteriza-se como aplicada, pois parte de um problema operacional real: a fragilidade de controles manuais, planilhas dispersas e comunicação informal em processos de atendimento e manutenção. A fundamentação teórica relaciona sistemas de informação, gestão de processos, inovação tecnológica, gestão de serviços e engenharia de software, sempre com referências diretas aos autores consultados. O desenvolvimento utilizou Laravel, React, TypeScript, Vite, PostgreSQL, Laravel Sanctum, Docker e serviços de e-mail transacional, com documentação técnica organizada e testes automatizados para regras críticas. A etapa mais recente incluiu ajustes de produção para Railway em 02 de junho de 2026, contemplando variáveis de banco PostgreSQL da plataforma, porta dinâmica de execução, preparação do Apache, endpoint de saúde e configuração voltada ao uso de domínio público com HTTPS. Como resultado, o TechOS Flow demonstra viabilidade técnica para centralizar dados operacionais, melhorar rastreabilidade, reduzir retrabalho e oferecer suporte gerencial por indicadores e relatórios.",
    )
    p = add_paragraph(
        doc,
        "Palavras-chave: sistema web; ordem de serviço; saneamento; gestão de processos; Railway; domínio; Laravel; React.",
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_page_break()

    add_paragraph(doc, "SUMÁRIO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    summary = [
        "1 INTRODUÇÃO",
        "2 FUNDAMENTAÇÃO TEÓRICA",
        "3 METODOLOGIA",
        "4 DESENVOLVIMENTO DO SISTEMA",
        "5 IMPLANTAÇÃO EM NUVEM, RAILWAY E DOMÍNIO",
        "6 RESULTADOS ALCANÇADOS E AVALIAÇÃO",
        "7 CONSIDERAÇÕES FINAIS",
        "REFERÊNCIAS",
        "APÊNDICE A - Linha do tempo técnica",
        "APÊNDICE B - Síntese dos requisitos implementados",
    ]
    add_bullets(doc, summary)
    doc.add_page_break()

    # 1
    doc.add_heading("1 INTRODUÇÃO", level=1)
    for text in [
        "A gestão de ordens de serviço é uma atividade central para organizações que dependem de operações contínuas, equipes de campo, controle de prazos e registro confiável de atendimento. No setor de saneamento, esse tipo de processo envolve solicitações relacionadas a ligações, reparos, manutenção de redes, acompanhamento de estações e registros técnicos que precisam ser compreendidos por atendentes, técnicos e gestores. Quando esse fluxo é apoiado apenas por formulários físicos, planilhas isoladas ou mensagens informais, surgem problemas de rastreabilidade, duplicidade de dados, perda de evidências, dificuldade de auditoria e baixa visibilidade gerencial.",
        "O TechOS Flow foi concebido para enfrentar esse cenário por meio de uma plataforma web que centraliza a abertura, o acompanhamento, a execução e a finalização de ordens de serviço. A solução evoluiu ao longo do projeto e passou a abranger não apenas o cadastro básico de OS, mas também autenticação segura, separação por perfis, dashboards, relatórios, anexos privados, geolocalização de evidências, gestão de usuários, colaboradores operacionais, horas extras, banco de folgas e exportações administrativas.",
        "A inovação proposta não se limita à troca do papel por uma tela. Ela reorganiza o fluxo informacional da operação, transformando registros dispersos em dados estruturados que podem apoiar decisões. Essa perspectiva dialoga com Kroenke (2012, p. 28), ao tratar sistemas de informação como combinações de hardware, software, dados, procedimentos e pessoas. No TechOS Flow, esses componentes aparecem de forma integrada: o navegador e o servidor compõem a infraestrutura, o software organiza regras e interfaces, o banco preserva dados, os procedimentos definem o ciclo da OS e as pessoas assumem papéis operacionais distintos.",
        "O problema de pesquisa que orienta este trabalho é: como um sistema web de gestão de ordens de serviço pode otimizar o controle operacional, reduzir falhas de comunicação interna e fornecer suporte gerencial em operações de saneamento?",
    ]:
        add_body(doc, text)

    doc.add_heading("1.1 Objetivo geral", level=2)
    add_body(
        doc,
        "Desenvolver, documentar e implantar em ambiente de nuvem um sistema web para automação do fluxo de ordens de serviço, desde a abertura até o encerramento, incluindo rastreabilidade, evidências, relatórios, segurança e suporte gerencial.",
    )
    doc.add_heading("1.2 Objetivos específicos", level=2)
    add_bullets(
        doc,
        [
            "mapear o processo operacional de geração, aceite, execução e encerramento de ordens de serviço;",
            "estruturar perfis de acesso para administrador, atendente e técnico;",
            "padronizar campos operacionais, endereços, status, prioridades e evidências;",
            "implementar relatórios, dashboards e exportações em PDF, Excel e CSV;",
            "registrar anexos privados com possibilidade de geolocalização em smartphone;",
            "apurar horas extras e banco de folgas por participante da execução;",
            "preparar o sistema para produção com Docker, Railway, PostgreSQL em nuvem, domínio público e HTTPS;",
            "documentar requisitos, arquitetura, banco de dados, segurança, testes e implantação.",
        ],
    )
    doc.add_heading("1.3 Estrutura do trabalho", level=2)
    add_body(
        doc,
        "O trabalho está organizado em fundamentação teórica, metodologia, desenvolvimento do sistema, implantação em nuvem, resultados alcançados, considerações finais, referências e apêndices técnicos. A estrutura preserva o formato do artigo original, mas atualiza o conteúdo conforme o estado real do projeto até a subida para Railway em 02 de junho de 2026.",
    )

    # 2
    doc.add_heading("2 FUNDAMENTAÇÃO TEÓRICA", level=1)
    doc.add_heading("2.1 Sistemas de informação", level=2)
    for text in [
        "Sistemas de informação são relevantes porque articulam tecnologia, procedimentos e pessoas para produzir dados úteis à organização. Kroenke (2012, p. 28) apresenta a estrutura de cinco componentes de um sistema de informação: hardware, software, dados, procedimentos e pessoas. Essa estrutura ajuda a compreender que a qualidade de um sistema não depende apenas do código, mas da forma como os usuários operam, interpretam e alimentam o fluxo informacional.",
        "No contexto do TechOS Flow, os dados de ordens de serviço, usuários, endereços, execuções, anexos e horas extras deixam de permanecer em registros fragmentados e passam a compor um repositório relacional. Kroenke (2012, p. 43) ressalta que repositórios concentram a memória operacional da empresa, podendo assumir desde listas simples até bancos de dados. Ao utilizar PostgreSQL, o projeto busca reduzir custos de oportunidade associados à perda de informação e à dificuldade de consulta.",
        "A qualidade da informação também depende de relevância, completude e adequação ao custo. Kroenke (2012, p. 46-47) relaciona informação útil a características como precisão, pontualidade, relevância e suficiência. Esse princípio orientou a criação dos dashboards por perfil e dos relatórios administrativos: o atendente precisa de visão operacional de abertura e acompanhamento, o técnico precisa de fila e execução, e o administrador precisa de indicadores agregados para decisão.",
    ]:
        add_body(doc, text)

    doc.add_heading("2.2 Gestão de processos", level=2)
    for text in [
        "A gestão de processos permite analisar atividades, papéis, recursos, dados e decisões envolvidos em uma função de negócio. Kroenke (2012, p. 41) define processo de negócio como uma rede de atividades, funções, recursos, repositórios e fluxos de dados que interagem para executar uma função organizacional. No caso do saneamento, a ordem de serviço representa um fluxo de atendimento que nasce em uma solicitação, passa por triagem, execução técnica, comprovação e encerramento.",
        "Oliveira (2019, p. 152) contribui para essa leitura ao tratar a administração de processos como uma forma de organizar atividades de modo integrado. A digitalização proposta pelo TechOS Flow não apenas registra etapas, mas cria critérios de passagem entre estados, como OS aberta, aceita, em execução, finalizada ou não executada. Isso reduz ambiguidade e cria uma sequência operacional verificável.",
        "Ao mapear processos, também se torna possível identificar gargalos, retrabalho e necessidades de controle. Oliveira (2019, p. 236; p. 260) reforça a importância de padronizar e melhorar processos para ampliar desempenho. No sistema desenvolvido, essa padronização aparece nas regras de negócio, nos status, nos campos obrigatórios, no vínculo entre OS e endereço, no registro de equipe e na exigência de motivo quando uma OS não é executada.",
    ]:
        add_body(doc, text)

    doc.add_heading("2.3 Inovação tecnológica e serviços", level=2)
    for text in [
        "A inovação tecnológica pode ser compreendida como adoção de novos métodos, produtos ou processos capazes de gerar melhoria organizacional. Andreassi (2012, p. 7) relaciona inovação à aplicação de mudanças que produzem valor. No TechOS Flow, o valor esperado surge da redução de papel, da rastreabilidade, da velocidade de consulta, do registro de evidências e da base de dados preparada para análise.",
        "A inovação também possui relação com desempenho e vantagem competitiva. Andreassi (2012, p. 13) destaca a importância da capacidade de inovar para gerar crescimento e diferenciação. Ainda que o projeto tenha caráter acadêmico e operacional, sua aplicação em saneamento pode contribuir para maior eficiência no atendimento ao cidadão e melhor alocação de equipes.",
        "A gestão de serviços exige atenção à percepção do usuário, à confiabilidade do processo e à qualidade da entrega. Corrêa e Caon (2012, p. 150; p. 174) discutem a relação entre operações de serviços, satisfação e desempenho. A ordem de serviço é um instrumento direto dessa relação, pois registra a demanda, orienta a equipe e produz evidências da execução.",
        "Tidd e Bessant (2015, p. 55) tratam a inovação como processo a ser gerenciado, e não apenas como evento isolado. Essa abordagem se aproxima da evolução incremental do projeto, que passou por ciclos de autenticação, execução, relatórios, segurança, geolocalização, documentação, Docker e implantação em nuvem.",
    ]:
        add_body(doc, text)

    doc.add_heading("2.4 Engenharia de software", level=2)
    for text in [
        "A engenharia de software fornece métodos para transformar necessidades reais em sistemas verificáveis, manuteníveis e evolutivos. Pressman e Maxim (2021, p. 20) associam software de qualidade à aplicação disciplinada de processo, métodos e ferramentas. O TechOS Flow adotou essa lógica ao separar frontend, backend, banco, documentação, testes e scripts de implantação.",
        "Pressman e Maxim (2021, p. 26) também destacam a importância de processos adaptáveis. O desenvolvimento do sistema ocorreu de forma incremental, permitindo que funcionalidades fossem consolidadas, testadas e ampliadas. A linha do tempo do repositório demonstra esse percurso: login e UUID em fevereiro de 2026, fluxo de execução em fevereiro, dashboards em março, segurança e relatórios em abril e maio, documentação e implantação no fim de maio e ajustes para Railway em junho.",
        "No campo de qualidade, Pressman e Maxim (2021, p. 437; p. 441; p. 446) tratam testes, verificação e validação como partes essenciais do desenvolvimento. No projeto, isso se materializa em testes automatizados de backend para autenticação, fluxo técnico, recuperação de senha, relatórios e horas extras, além de validações de frontend como lint, TypeScript, build e acessibilidade.",
    ]:
        add_body(doc, text)

    doc.add_heading("2.5 Sistemas gerenciais e apoio à decisão", level=2)
    for text in [
        "O'Brien e Marakas (2012, p. 30; p. 34) associam sistemas de informação ao apoio às operações, à administração e à tomada de decisão. O TechOS Flow atende a essas três dimensões: apoia a operação com abertura e execução de OS, apoia a administração com usuários, colaboradores e horas extras, e apoia a decisão com dashboards e relatórios.",
        "A perspectiva de Business Intelligence também contribui para a justificativa do sistema. Kroenke (2012, p. 205) caracteriza sistemas de BI como sistemas que fornecem informações para melhorar o processo decisório. Embora o TechOS Flow não seja uma plataforma completa de BI, os relatórios administrativos e exportações estruturadas criam a base para indicadores operacionais, como quantidade de OS por status, produtividade por técnico, horas extras e distribuição por prioridade.",
    ]:
        add_body(doc, text)

    # 3
    doc.add_heading("3 METODOLOGIA", level=1)
    for text in [
        "A metodologia adotada é aplicada, documental e construtiva. É aplicada porque busca resolver um problema prático de gestão operacional de ordens de serviço. É documental porque utiliza modelos de OS, documentação técnica, requisitos e registros do repositório. É construtiva porque resulta em um artefato de software funcional, acompanhado de documentação e validações.",
        "O desenvolvimento foi conduzido de modo incremental. Cada ciclo agregou funcionalidades e revisões: autenticação, cadastro de endereços, ciclo de execução, dashboards, relatórios, evidências, geolocalização, horas extras, banco de folgas, segurança, documentação, Docker e implantação. Essa estratégia permitiu incorporar ajustes reais de uso, corrigir limitações e preparar o sistema para ambiente de nuvem.",
        "As fontes internas utilizadas nesta atualização foram o README do projeto, documentos de visão, requisitos, arquitetura, banco de dados, segurança, plano de testes, implantação, histórico de commits e arquivos de produção. As fontes bibliográficas foram usadas diretamente, por autor, ano e página.",
    ]:
        add_body(doc, text)
    doc.add_heading("3.1 Etapas metodológicas", level=2)
    add_table(
        doc,
        ["Etapa", "Descrição", "Produto gerado"],
        [
            ("Levantamento", "Análise do fluxo de OS, perfis e dados necessários.", "Problema, objetivos e requisitos."),
            ("Modelagem", "Definição de entidades, casos de uso, status e regras.", "Documentos de requisitos, banco e arquitetura."),
            ("Implementação", "Construção do backend Laravel, frontend React e banco PostgreSQL.", "Sistema web funcional."),
            ("Validação", "Execução de testes automatizados e inspeções manuais por perfil.", "Plano de testes e critérios de aceite."),
            ("Implantação", "Preparação de Docker, Railway, variáveis, health check e domínio/HTTPS.", "Ambiente publicável em nuvem."),
            ("Documentação", "Consolidação de manuais, API, segurança, implantação e TCC.", "Base acadêmica e técnica atualizada."),
        ],
        [3.0, 8.0, 5.0],
    )

    # 4
    doc.add_heading("4 DESENVOLVIMENTO DO SISTEMA", level=1)
    doc.add_heading("4.1 Visão geral da solução", level=2)
    for text in [
        "O TechOS Flow é um sistema web para gestão digital de ordens de serviço. A aplicação possui frontend em React, Vite e TypeScript, backend em Laravel, API REST versionada em /api/v1, autenticação com Laravel Sanctum e banco PostgreSQL. A arquitetura separa apresentação, aplicação, domínio, persistência e infraestrutura, facilitando manutenção e evolução.",
        "Os perfis implementados são administrador, atendente e técnico. O atendente pode criar e acompanhar ordens gerais; o técnico pode aceitar, iniciar, finalizar, não executar e criar OS de Manutenção ETA/ETE; o administrador acompanha relatórios, usuários, colaboradores operacionais, PDFs, exportações e horas extras.",
    ]:
        add_body(doc, text)
    doc.add_heading("4.2 Requisitos funcionais implementados", level=2)
    add_table(
        doc,
        ["Código", "Requisito atualizado", "Situação"],
        [
            ("RF-01", "Autenticar usuários por e-mail e senha, com logout e consulta do usuário autenticado.", "Implementado"),
            ("RF-02", "Exigir troca de senha forte no primeiro acesso.", "Implementado"),
            ("RF-03", "Permitir recuperação de senha por e-mail transacional.", "Implementado"),
            ("RF-04", "Criar OS geral por atendente e OS Manutenção ETA/ETE por técnico.", "Implementado"),
            ("RF-05", "Listar, filtrar, paginar e detalhar ordens de serviço.", "Implementado"),
            ("RF-06", "Permitir aceite, início, finalização e não execução da OS pelo técnico responsável.", "Implementado"),
            ("RF-07", "Permitir anexos e evidências com geolocalização opcional.", "Implementado"),
            ("RF-08", "Compor equipe com usuários e colaboradores operacionais sem login.", "Implementado"),
            ("RF-09", "Calcular horas extras 50%, 100% e banco de folgas por participante.", "Implementado"),
            ("RF-10", "Gerar dashboards, relatórios, PDF detalhado, Excel e CSV.", "Implementado"),
            ("RF-11", "Gerenciar usuários, inativação e colaboradores operacionais.", "Implementado"),
            ("RF-12", "Manter anexos privados com acesso autenticado e autorizado.", "Implementado"),
        ],
        [2.0, 11.0, 3.0],
    )
    doc.add_heading("4.3 Requisitos não funcionais e qualidade", level=2)
    add_bullets(
        doc,
        [
            "Arquitetura web em camadas, com separação entre frontend, backend e banco.",
            "API versionada em /api/v1 e autenticação baseada em Sanctum.",
            "Interface responsiva para desktop e smartphone.",
            "Uso de UUID nas entidades centrais.",
            "Relatórios e exportações calculados no backend para reduzir duplicação de regra.",
            "Armazenamento privado de anexos.",
            "Validação de acessibilidade, lint, tipagem e build no frontend.",
            "Testes automatizados de backend para regras críticas.",
        ],
    )
    doc.add_heading("4.4 Banco de dados", level=2)
    add_body(
        doc,
        "O banco PostgreSQL foi modelado para preservar a rastreabilidade da operação. As principais tabelas são users, enderecos, ordem_servicos, execucoes, execucao_funcionarios, colaboradores_operacionais, anexos, feriados, password_reset_tokens, personal_access_tokens e sessions. A separação entre usuários autenticáveis e colaboradores operacionais evita conceder login a participantes que apenas compõem a equipe de execução, mantendo o controle administrativo mais simples e seguro.",
    )
    add_table(
        doc,
        ["Entidade", "Finalidade"],
        [
            ("users", "Usuários autenticáveis, papéis, status e valor-hora."),
            ("ordem_servicos", "Núcleo do atendimento, status, prioridade, descrição e vínculo com endereço."),
            ("execucoes", "Histórico de início, fim e observações da execução técnica."),
            ("execucao_funcionarios", "Participantes da execução e base de cálculo de horas extras."),
            ("colaboradores_operacionais", "Participantes sem login, usados em equipe e horas extras."),
            ("anexos", "Evidências, arquivos, coordenadas e endereço capturado em campo."),
            ("feriados", "Apoio ao cálculo de horas extras e banco de folgas."),
        ],
        [4.0, 12.0],
    )
    doc.add_heading("4.5 Segurança, privacidade e LGPD", level=2)
    for text in [
        "A segurança foi tratada como requisito transversal. O sistema implementa login por e-mail e senha, tokens Sanctum, bloqueio de usuários inativos, primeiro acesso com troca obrigatória de senha, recuperação por token temporário, controle por perfil e restrição de acesso a anexos.",
        "Os dados tratados incluem dados de usuários, dados operacionais da OS, endereço, execução, evidências, geolocalização, sessão e recuperação de senha. Como evidências e coordenadas podem conter informações sensíveis do contexto operacional, o projeto adota anexos privados e acesso mediado pelo backend. A documentação também prevê retenção assistida de anexos por comando administrativo e complementação futura de política institucional de descarte.",
    ]:
        add_body(doc, text)
    doc.add_heading("4.6 Interface e usabilidade", level=2)
    add_body(
        doc,
        "A interface foi organizada por módulos e perfis. O usuário acessa painéis específicos e executa fluxos compatíveis com suas permissões. Foram adicionadas melhorias de responsividade, tema claro/escuro, mensagens de erro em português, filtros, paginação, miniaturas de fotos nos detalhes da OS e indicação de sucesso ou falha na captura de geolocalização.",
    )
    doc.add_heading("4.7 Testes e validação", level=2)
    add_body(
        doc,
        "A estratégia de testes combina validações automatizadas e manuais. No backend, há testes para autenticação e segurança, dashboard, fluxo técnico, primeiro acesso, recuperação de senha, relatórios administrativos, retenção de anexos, gestão de usuários, horas extras e colaboradores operacionais. No frontend, há lint, TypeScript, build, testes de utilitários e suíte de acessibilidade. A regressão mínima recomendada inclui php artisan test, npm run lint, npx tsc -b, npm run build e npm run test:a11y.",
    )

    # 5
    doc.add_heading("5 IMPLANTAÇÃO EM NUVEM, RAILWAY E DOMÍNIO", level=1)
    for text in [
        "A etapa final do projeto até esta atualização concentrou-se na preparação para execução em ambiente real. Antes da escolha prática do Railway, o repositório recebeu uma pilha de produção para Oracle Cloud com Caddy, Docker, PostgreSQL, volumes persistentes e HTTPS. Essa etapa serviu como base de compreensão dos requisitos de produção: domínio, certificados, banco persistente, storage para anexos, SMTP, backup e validação em smartphone.",
        "Em seguida, o projeto foi ajustado para Railway. O histórico de commits mostra que, em 01 de junho de 2026, foi removido conflito de módulos MPM do Apache na imagem de backend. Em 02 de junho de 2026, foram realizados ajustes sucessivos para porta dinâmica do Railway, inicialização de produção, variáveis PostgreSQL da plataforma, endpoint de saúde na raiz do backend e commit final para disparar o deploy.",
        "A imagem de produção do backend passou a usar PHP 8.2 com Apache, extensões para PostgreSQL, ZIP, GD e OPcache, Composer otimizado, script de inicialização e exposição da porta 8080. O script start-production configura a porta a partir da variável PORT, converte variáveis PGHOST, PGPORT, PGDATABASE, PGUSER e PGPASSWORD para o padrão esperado pelo Laravel, aguarda o PostgreSQL ficar disponível, executa migrations quando autorizado, limpa e gera caches de configuração, rotas e views, e inicia o Apache.",
        "No frontend, a imagem de produção usa Node 20 para gerar o build do Vite e Nginx para servir os arquivos estáticos. A variável VITE_API_URL foi preparada para operar como /api/v1 quando frontend e backend são publicados sob o mesmo domínio ou atrás de proxy, reduzindo problemas de CORS e simplificando a experiência do usuário.",
        "O uso de domínio público foi incorporado ao desenho de implantação por variáveis como APP_URL, FRONTEND_URL e SANCTUM_STATEFUL_DOMAINS. A documentação do projeto usa techosflow.com.br como domínio de referência, com api.techosflow.com.br para API, além do uso local techosflow.test em desenvolvimento. Em produção, o domínio com HTTPS é essencial para proteger credenciais, preservar sessões, permitir recuperação de senha segura e viabilizar geolocalização confiável no navegador do smartphone.",
    ]:
        add_body(doc, text)
    doc.add_heading("5.1 Linha do tempo de implantação", level=2)
    add_table(
        doc,
        ["Data", "Marco técnico", "Impacto"],
        [
            ("31/05/2026", "Adição de pilha de produção para Oracle Cloud e smoke test local.", "Base para implantação com domínio, HTTPS e volumes."),
            ("01/06/2026", "Desativação de módulos Apache MPM conflitantes na imagem de backend.", "Correção de incompatibilidade de runtime."),
            ("02/06/2026", "Correção do binding de porta Apache para Railway.", "Adequação à porta dinâmica da plataforma."),
            ("02/06/2026", "Ajuste do start-production e suporte às variáveis PostgreSQL do Railway.", "Conexão com banco gerenciado da plataforma."),
            ("02/06/2026", "Endpoint de saúde na raiz do backend.", "Facilita verificação de disponibilidade."),
            ("02/06/2026", "Commit Trigger Railway deploy.", "Subida final para acionar implantação."),
        ],
        [3.0, 7.0, 6.0],
    )
    doc.add_heading("5.2 Variáveis e cuidados de produção", level=2)
    add_bullets(
        doc,
        [
            "APP_ENV=production e APP_DEBUG=false para ambiente publicado.",
            "APP_URL e FRONTEND_URL alinhados ao domínio público.",
            "SANCTUM_STATEFUL_DOMAINS configurado conforme o domínio usado.",
            "DB_HOST, DB_PORT, DB_DATABASE, DB_USERNAME e DB_PASSWORD derivados das variáveis PostgreSQL do Railway.",
            "MAIL_* configurado para SMTP real, com referência ao Zoho Mail no projeto.",
            "RUN_MIGRATIONS_ON_START controlando execução de migrations em produção.",
            "Storage persistente para anexos, necessário para que evidências não sejam perdidas em reinícios.",
        ],
    )

    # 6
    doc.add_heading("6 RESULTADOS ALCANÇADOS E AVALIAÇÃO", level=1)
    for text in [
        "A versão atual do TechOS Flow apresenta avanço significativo em relação à proposta inicial. O sistema deixou de ser um MVP conceitual e passou a reunir funcionalidades de operação, gestão, segurança, documentação, testes e implantação. O fluxo de OS está completo para abertura, aceite, execução, finalização e não execução, com regras por perfil e controle de acesso.",
        "O resultado operacional mais relevante é a rastreabilidade. Cada OS possui número, status, prioridade, endereço, responsável, histórico de execução, equipe e anexos. Isso melhora a capacidade de consulta e reduz dependência de registros manuais. O resultado gerencial está nos dashboards, relatórios, PDFs e exportações, que permitem acompanhar volume de trabalho, produtividade e horas extras.",
        "A etapa de evidências adicionou valor ao uso em campo, pois permite anexar fotos e capturar geolocalização. Essa funcionalidade se conecta diretamente ao requisito de domínio com HTTPS, já que navegadores modernos restringem geolocalização em contextos inseguros. A implantação em Railway e o uso planejado de domínio público representam, portanto, não apenas uma escolha de hospedagem, mas parte da validação funcional do sistema.",
    ]:
        add_body(doc, text)
    doc.add_heading("6.1 Indicadores sugeridos", level=2)
    add_table(
        doc,
        ["Indicador", "Forma de medição", "Uso gerencial"],
        [
            ("OS por status", "Contagem de abertas, em execução, finalizadas e não executadas.", "Acompanhar fila e gargalos."),
            ("Tempo médio de atendimento", "Diferença entre abertura e encerramento.", "Avaliar eficiência operacional."),
            ("OS com evidência", "Percentual de OS finalizadas com anexos.", "Acompanhar qualidade de comprovação."),
            ("Produtividade por técnico", "OS finalizadas por período e técnico.", "Distribuir carga e identificar padrões."),
            ("Horas extras", "Minutos 50%, 100% e banco de folgas.", "Controlar custo e compensação."),
            ("Prioridades", "Distribuição por nível de prioridade.", "Planejar atendimento e urgências."),
        ],
        [4.0, 6.0, 6.0],
    )
    doc.add_heading("6.2 Riscos e limitações", level=2)
    add_bullets(
        doc,
        [
            "Ainda não há aplicativo mobile nativo; o uso em smartphone ocorre pelo navegador.",
            "A persistência de anexos em produção depende de volume ou storage configurado corretamente.",
            "A validação completa de geolocalização depende de domínio final com HTTPS.",
            "Observabilidade, monitoramento e alertas ainda estão em nível inicial.",
            "Políticas institucionais de LGPD, backup, descarte e incidente precisam ser formalizadas fora do código.",
            "Testes E2E e testes de carga são melhorias futuras recomendadas.",
        ],
    )

    # 7
    doc.add_heading("7 CONSIDERAÇÕES FINAIS", level=1)
    for text in [
        "O desenvolvimento do TechOS Flow demonstra a viabilidade de um sistema web para organizar a gestão de ordens de serviço no setor de saneamento. A solução centraliza dados, padroniza o fluxo de trabalho, diferencia perfis, registra evidências e fornece relatórios gerenciais. Com isso, responde ao problema de pesquisa ao mostrar que a digitalização estruturada pode reduzir falhas de comunicação, ampliar rastreabilidade e apoiar decisões administrativas.",
        "A fundamentação teórica mostrou que sistemas de informação dependem da integração entre tecnologia, dados, procedimentos e pessoas; que processos de negócio precisam ser mapeados e melhorados; que a inovação tecnológica deve gerar valor organizacional; e que engenharia de software exige processo, qualidade e validação. Esses elementos aparecem de forma concreta no projeto implementado.",
        "A atualização mais importante em relação ao texto original está na chegada do projeto a uma etapa de implantação. A preparação para Railway em 02 de junho de 2026, associada ao uso de domínio público e HTTPS, aproxima o sistema de uma validação real. Essa etapa é decisiva para testar recuperação de senha, geolocalização, anexos, relatórios e estabilidade fora do ambiente local.",
        "Como trabalhos futuros, recomenda-se validar o sistema com usuários reais, configurar storage persistente definitivo, automatizar backups, ampliar monitoramento, criar testes E2E, avaliar aplicativo mobile híbrido ou nativo e formalizar políticas institucionais de privacidade e retenção. Também é recomendável evoluir os indicadores para análises comparativas e séries históricas, aproximando o sistema de uma camada mais robusta de apoio à decisão.",
    ]:
        add_body(doc, text)

    # References
    doc.add_heading("REFERÊNCIAS", level=1)
    refs = [
        "ANDREASSI, Tales. Gestão da Inovação Tecnológica. Porto Alegre: +A Educação - Cengage Learning Brasil, 2012. E-book. ISBN 9788522108404. Disponível em: https://integrada.minhabiblioteca.com.br/reader/books/9788522108404/. Acesso em: 01 nov. 2025.",
        "CORRÊA, Henrique L.; CAON, Mauro. Gestão de serviços: lucratividade por meio de operações e de satisfação dos clientes. Rio de Janeiro: Atlas, 2012. E-book. ISBN 9788522479214. Disponível em: https://integrada.minhabiblioteca.com.br/reader/books/9788522479214/. Acesso em: 01 nov. 2025.",
        "FILHO, Wilson de Pádua P. Engenharia de Software: projetos e processos. Vol. 2. 4. ed. Rio de Janeiro: LTC, 2019. E-book. ISBN 9788521636748. Disponível em: https://integrada.minhabiblioteca.com.br/reader/books/9788521636748/. Acesso em: 01 nov. 2025.",
        "KROENKE, David M. Sistemas de informação gerenciais. 1. ed. Rio de Janeiro: Saraiva, 2012. E-book. ISBN 9788502183704. Disponível em: https://integrada.minhabiblioteca.com.br/reader/books/9788502183704/. Acesso em: 31 out. 2025.",
        "O'BRIEN, James A.; MARAKAS, George M. Administração de sistemas de informação. 15. ed. Porto Alegre: AMGH, 2012. E-book. ISBN 9788580551112. Disponível em: https://integrada.minhabiblioteca.com.br/reader/books/9788580551112/. Acesso em: 02 nov. 2025.",
        "OLIVEIRA, Djalma de Pinho Rebouças de. Administração de Processos. 6. ed. Rio de Janeiro: Atlas, 2019. E-book. ISBN 9788597021301. Disponível em: https://integrada.minhabiblioteca.com.br/reader/books/9788597021301/. Acesso em: 01 nov. 2025.",
        "PRESSMAN, Roger S.; MAXIM, Bruce R. Engenharia de software. 9. ed. Porto Alegre: AMGH, 2021. E-book. ISBN 9786558040118. Disponível em: https://integrada.minhabiblioteca.com.br/reader/books/9786558040118/. Acesso em: 02 nov. 2025.",
        "TIDD, Joe; BESSANT, Joe. Gestão da inovação. 5. ed. Porto Alegre: Bookman, 2015. E-book. ISBN 9788582603079. Disponível em: https://integrada.minhabiblioteca.com.br/reader/books/9788582603079/. Acesso em: 02 nov. 2025.",
    ]
    for ref in refs:
        p = doc.add_paragraph(ref)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)

    # Appendices
    doc.add_heading("APÊNDICE A - Linha do tempo técnica", level=1)
    add_table(
        doc,
        ["Data", "Evolução do projeto"],
        [
            ("08/02/2026", "Primeiros commits e estrutura inicial do repositório."),
            ("15/02/2026", "Login, logout, consulta do usuário autenticado, Sanctum e UUID funcionando."),
            ("16/02/2026", "Fluxo completo de execução finalizado com ciclo de status."),
            ("17/02/2026", "Autenticação, endereço, geocoding, ordem de serviço e execução finalizados."),
            ("19/03/2026", "Ajuste de dashboard por perfil e fluxo interno de OS."),
            ("08/04/2026", "Conclusão de blocos operacional e gerencial, primeiro acesso e melhorias de relatórios."),
            ("16/05/2026", "Consolidação de funcionalidades, segurança e documentação do projeto."),
            ("30/05/2026", "Colaboradores operacionais, evidências, relatórios, exportações, responsividade e recuperação de senha."),
            ("31/05/2026", "Pilha de produção para Oracle Cloud e smoke test local."),
            ("01/06/2026", "Correção de módulos Apache para imagem de produção."),
            ("02/06/2026", "Ajustes finais para Railway, PostgreSQL, porta dinâmica, health check e deploy."),
        ],
        [3.0, 13.0],
    )

    doc.add_heading("APÊNDICE B - Síntese dos requisitos implementados", level=1)
    add_bullets(
        doc,
        [
            "Autenticação por e-mail e senha, logout, usuário autenticado, primeiro acesso e recuperação por e-mail.",
            "Perfis administrador, atendente e técnico, com autorização de rotas e ações.",
            "Criação de OS geral, criação de OS ETA/ETE, filtros, paginação, detalhes e número único.",
            "Aceite, início, finalização e não execução de OS com validações no backend.",
            "Evidências com anexos privados, coordenadas e endereço capturado quando disponível.",
            "Equipe de execução com usuários e colaboradores operacionais.",
            "Horas extras, banco de folgas, feriados e relatórios administrativos.",
            "Exportação PDF, Excel e CSV, além de PDF detalhado da OS restrito ao administrador.",
            "Docker de desenvolvimento, Docker de produção, ajustes para Railway e configuração de domínio/HTTPS.",
        ],
    )

    # Metadata
    doc.core_properties.title = "TechOS Flow: Sistema Web para Gestão Digital de Ordens de Serviço no Setor de Saneamento"
    doc.core_properties.author = "Mateus Lima da Silva Braga"
    doc.core_properties.subject = "TCC atualizado com implantação Railway e domínio"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
